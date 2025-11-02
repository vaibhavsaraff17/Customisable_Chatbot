from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
import requests
import os
import json
import uuid
import shutil
from datetime import datetime
from pathlib import Path
from pymongo import MongoClient
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from werkzeug.utils import secure_filename
from docx import Document as DocxDocument
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configuration
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434/api/generate")
MODEL_NAME = os.getenv("MODEL_NAME", "llama3.2:3b")
UPLOAD_FOLDER = os.getenv("UPLOAD_FOLDER", "vector_stores")
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'json'}
DEFAULT_SYSTEM_PROMPT = (
    "You are a smart assistant that strictly follows the user's custom instructions.\n"
    "IMPORTANT: You must ONLY answer questions based on the content provided in the 'Relevant Information' section below. "
    "This information comes from uploaded documents and is the ONLY source you should use for answers.\n"
    "Do NOT use your general knowledge or training data unless explicitly instructed otherwise through custom instructions.\n"
    "Never include or mention document names, file paths, or any metadata in your answers.\n"
    "If a user's question is vague, incomplete, or ambiguous, ask follow-up questions for clarification "
    "instead of making assumptions or guessing.\n"
    "If the answer is not present in the provided 'Relevant Information' section and no custom instruction allows outside knowledge, "
    "clearly state: 'I cannot find this information in the uploaded documents.'\n"
    "Encourage a multi-turn conversation by asking relevant follow-up questions that ensure clarity and help guide the user.\n"
    "When providing answers, format them clearly using bullet points, numbered steps, or concise paragraphs for better readability and understanding.\n"
    "Highlight important concepts, keywords, or section titles using **bold Markdown formatting** to improve clarity and emphasis.\n"

)


# Flask setup
app = Flask(__name__)
CORS(app)  # Enable CORS for all routes
app.config['DEBUG'] = True
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size

# MongoDB connection - Load from environment variable
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/")

try:
    client = MongoClient(MONGO_URI, serverSelectionTimeoutMS=5000)
    # Test the connection
    client.admin.command('ping')
    print("MongoDB connection successful!")
    db = client["enhanced_chatbot"]
    sessions_collection = db["sessions"]
    conversations_collection = db["conversations"]
except Exception as e:
    print(f"MongoDB connection failed: {e}")
    print("Falling back to local file storage...")
    
    # Fallback to local JSON storage
    import json
    from pathlib import Path
    
    LOCAL_DB_PATH = Path("local_db")
    LOCAL_DB_PATH.mkdir(exist_ok=True)
    
    class LocalQuery:
        def __init__(self, data, query=None):
            self.data = data
            self.query = query or {}
            self._sort_field = None
            self._sort_order = 1
            self._limit_count = None
        
        def sort(self, field, order=1):
            self._sort_field = field
            self._sort_order = order
            return self
        
        def limit(self, count):
            self._limit_count = count
            return self
        
        def __iter__(self):
            # Filter data based on query
            filtered_data = []
            for doc in self.data:
                if all(doc.get(k) == v for k, v in self.query.items()):
                    filtered_data.append(doc)
            
            # Sort if specified
            if self._sort_field:
                filtered_data.sort(
                    key=lambda x: x.get(self._sort_field, ''),
                    reverse=(self._sort_order == -1)
                )
            
            # Apply limit if specified
            if self._limit_count:
                filtered_data = filtered_data[:self._limit_count]
            
            return iter(filtered_data)
    
    class LocalCollection:
        def __init__(self, name):
            self.file_path = LOCAL_DB_PATH / f"{name}.json"
            self.data = self._load_data()
        
        def _load_data(self):
            if self.file_path.exists():
                with open(self.file_path, 'r') as f:
                    return json.load(f)
            return []
        
        def _save_data(self):
            with open(self.file_path, 'w') as f:
                json.dump(self.data, f, indent=2, default=str)
        
        def insert_one(self, doc):
            doc['_id'] = str(uuid.uuid4())
            self.data.append(doc)
            self._save_data()
            return doc
        
        def find_one(self, query):
            for doc in self.data:
                if all(doc.get(k) == v for k, v in query.items()):
                    return doc
            return None
        
        def find(self, query=None, projection=None):
            return LocalQuery(self.data, query)
        
        def update_one(self, query, update):
            for doc in self.data:
                if all(doc.get(k) == v for k, v in query.items()):
                    if '$set' in update:
                        doc.update(update['$set'])
                    self._save_data()
                    return doc
            return None
        
        def delete_many(self, query):
            original_count = len(self.data)
            self.data = [doc for doc in self.data if not all(doc.get(k) == v for k, v in query.items())]
            deleted_count = original_count - len(self.data)
            self._save_data()
            return type('Result', (), {'deleted_count': deleted_count})()
    
    class LocalDB:
        def __init__(self):
            pass
        
        def __getitem__(self, name):
            return LocalCollection(name)
    
    db = LocalDB()
    sessions_collection = db["sessions"]
    conversations_collection = db["conversations"]

# Ensure upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize embeddings model
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_session_path(session_id):
    return os.path.join(UPLOAD_FOLDER, session_id)

def get_documents_path(session_id):
    return os.path.join(get_session_path(session_id), "documents")

def get_vector_store_path(session_id):
    return os.path.join(get_session_path(session_id), "faiss_index")

def create_session_directories(session_id):
    session_path = get_session_path(session_id)
    documents_path = get_documents_path(session_id)
    vector_store_path = get_vector_store_path(session_id)
    
    os.makedirs(session_path, exist_ok=True)
    os.makedirs(documents_path, exist_ok=True)
    os.makedirs(vector_store_path, exist_ok=True)
    
    return session_path, documents_path, vector_store_path

def load_docx(file_path):
    """Load content from DOCX file"""
    try:
        docx_file = DocxDocument(file_path)
        full_text = "\n".join([para.text for para in docx_file.paragraphs])
        return [Document(page_content=full_text, metadata={"source": file_path})]
    except Exception as e:
        print(f"Error loading DOCX file {file_path}: {e}")
        return []

def load_json_file(file_path):
    """Load content from JSON file"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        if isinstance(data, dict):
            content = json.dumps(data, indent=2)
        elif isinstance(data, list):
            content = "\n".join([json.dumps(item, indent=2) for item in data])
        else:
            content = str(data)
            
        return [Document(page_content=content, metadata={"source": file_path})]
    except Exception as e:
        print(f"Error loading JSON file {file_path}: {e}")
        return []

def process_document(file_path):
    """Process a single document and return Document objects"""
    file_extension = Path(file_path).suffix.lower()
    
    try:
        if file_extension == '.txt':
            loader = TextLoader(file_path, encoding='utf-8')
            return loader.load()
        elif file_extension == '.pdf':
            loader = PyPDFLoader(file_path)
            return loader.load()
        elif file_extension == '.docx':
            return load_docx(file_path)
        elif file_extension == '.json':
            return load_json_file(file_path)
        else:
            print(f"Unsupported file type: {file_extension}")
            return []
    except Exception as e:
        print(f"Error processing document {file_path}: {e}")
        return []

def build_vector_store_for_session(session_id):
    """Build or rebuild vector store for a specific session"""
    documents_path = get_documents_path(session_id)
    vector_store_path = get_vector_store_path(session_id)
    
    # Load all documents in the session's documents folder
    all_docs = []
    for file_path in Path(documents_path).glob("*"):
        if file_path.is_file() and allowed_file(file_path.name):
            docs = process_document(str(file_path))
            all_docs.extend(docs)
    
    if not all_docs:
        print(f"No documents found for session {session_id}")
        return False
    
    # Split documents into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    chunks = text_splitter.split_documents(all_docs)
    
    # Create vector store
    try:
        vector_store = FAISS.from_documents(chunks, embeddings)
        vector_store.save_local(vector_store_path)
        print(f"Vector store built for session {session_id} with {len(chunks)} chunks")
        return True
    except Exception as e:
        print(f"Error building vector store for session {session_id}: {e}")
        return False

def load_vector_store_for_session(session_id):
    """Load vector store for a specific session"""
    vector_store_path = get_vector_store_path(session_id)
    
    try:
        if os.path.exists(vector_store_path) and os.listdir(vector_store_path):
            return FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
        else:
            return None
    except Exception as e:
        print(f"Error loading vector store for session {session_id}: {e}")
        return None

def retrieve_context_for_session(session_id, query, k=10):  # Increased to retrieve more relevant chunks
    """Retrieve context from session-specific vector store"""
    vector_store = load_vector_store_for_session(session_id)
    
    if vector_store is None:
        return ""
    
    try:
        results = vector_store.similarity_search(query, k=k)
        # No limit on context length since we removed num_ctx limit
        context_parts = []
        for doc in results:
            context_parts.append(doc.page_content)
        return "\n\n".join(context_parts)
    except Exception as e:
        print(f"Error retrieving context for session {session_id}: {e}")
        return ""

def get_conversation_context(session_id, conversation_id=None, limit=5):  # Increased to 5 exchanges
    """Get recent conversation context"""
    query = {"session_id": session_id}
    if conversation_id:
        query["conversation_id"] = conversation_id
    
    recent_messages = list(conversations_collection.find(query)
                          .sort("timestamp", -1)
                          .limit(limit * 2))  # Get more to account for user/bot pairs
    
    context = []
    for msg in reversed(recent_messages):
        if msg.get("message_type") == "user":
            # No truncation - include full messages
            message = msg.get('message', '')
            context.append(f"User: {message}")
        elif msg.get("message_type") == "bot":
            # No truncation - include full messages
            message = msg.get('message', '')
            context.append(f"Assistant: {message}")
    
    return "\n".join(context[-10:])  # Last 5 exchanges (10 messages)

def query_llm_with_session(session_id, query, conversation_id=None, custom_prompt=None):
    """Query LLM with session-specific context and custom prompt"""
    # Get document context
    doc_context = retrieve_context_for_session(session_id, query)
    
    # Get conversation context
    conv_context = get_conversation_context(session_id, conversation_id)
    
    # Get session configuration
    session_data = sessions_collection.find_one({"session_id": session_id})
    user_prompt = ""
    if session_data and session_data.get("custom_prompt"):
        user_prompt = session_data["custom_prompt"]
    elif custom_prompt:
        user_prompt = custom_prompt

    # Final system prompt = predefined + user
    if user_prompt:
        system_prompt = f"{DEFAULT_SYSTEM_PROMPT}\n\nCustom Instructions:\n{user_prompt}"
    else:
        system_prompt = DEFAULT_SYSTEM_PROMPT

    # Build the complete prompt
    print(system_prompt)
    prompt_parts = [system_prompt]
    
    if doc_context:
        prompt_parts.append(f"\nRelevant Information:\n{doc_context}")
    
    if conv_context:
        prompt_parts.append(f"\nConversation History:\n{conv_context}")
    
    prompt_parts.append(f"\nUser Question: {query}\n\nAssistant:")
    
    full_prompt = "\n".join(prompt_parts)
    
    # Query Ollama with extended timeout
    try:
        response = requests.post(OLLAMA_URL, json={
            "model": MODEL_NAME,
            "prompt": full_prompt,
            "stream": False,
            "options": {
                "temperature": 0.1,  # Very focused/deterministic responses
                "top_k": 40,
                "top_p": 0.8,  # More conservative word selection
                "num_ctx": 32768  # Large context window for llama3.2
            }
        }, timeout=120)  # Increased timeout for larger contexts
        
        if response.status_code == 200:
            return response.json().get("response", "No response received.")
        else:
            return f"**Ollama Error ({response.status_code})**: The local AI server returned an error. Please ensure Ollama is running with the `{MODEL_NAME}` model installed."
    except requests.exceptions.Timeout:
        return "**Timeout Error**: The AI model is taking longer than expected. This often happens on the first request when the model needs to load into memory. Please try again - subsequent requests should be faster."
    except requests.exceptions.ConnectionError:
        return f"**Connection Error**: Cannot connect to Ollama server at {OLLAMA_URL}. Please start Ollama by running `ollama serve` in your terminal, then ensure the `{MODEL_NAME}` model is installed with `ollama pull {MODEL_NAME}`."
    except Exception as e:
        return f"**AI Service Error**: {str(e)}"

# API Routes

@app.route("/")
def index():
    return render_template("dashboard.html")

@app.route("/api/health", methods=["GET"])
def health_check():
    """Check if Ollama is running and model is available"""
    try:
        # Test connection to Ollama
        response = requests.get("http://localhost:11434/api/tags", timeout=5)
        if response.status_code == 200:
            models = response.json().get("models", [])
            model_names = [model.get("name", "") for model in models]
            
            llama3_available = any(MODEL_NAME in name for name in model_names)
            
            return jsonify({
                "ollama_running": True,
                "model_available": llama3_available,
                "available_models": model_names,
                "target_model": MODEL_NAME,
                "message": "Ollama is running!" if llama3_available else f"Ollama is running but {MODEL_NAME} model not found. Run: ollama pull {MODEL_NAME}"
            })
        else:
            return jsonify({
                "ollama_running": False,
                "model_available": False,
                "message": "Ollama server responded with error"
            }), 503
    except requests.exceptions.ConnectionError:
        return jsonify({
            "ollama_running": False,
            "model_available": False,
            "message": "Ollama is not running. Start it with: ollama serve"
        }), 503
    except Exception as e:
        return jsonify({
            "ollama_running": False,
            "model_available": False,
            "message": f"Health check failed: {str(e)}"
        }), 503

@app.route("/api/sessions/create", methods=["POST"])
def create_session():
    """Create a new user session"""
    data = request.get_json() or {}
    
    session_id = str(uuid.uuid4())
    user_description = data.get("user_description", "")
    use_case = data.get("use_case", "")
    
    # Create session directories
    create_session_directories(session_id)
    
    # Store session in database
    session_doc = {
        "session_id": session_id,
        "user_description": user_description,
        "use_case": use_case,
        "created_at": datetime.utcnow(),
        "custom_prompt": "",
        "documents_count": 0
    }
    
    sessions_collection.insert_one(session_doc)
    
    return jsonify({
        "session_id": session_id,
        "created_at": session_doc["created_at"].isoformat()
    })

@app.route("/api/sessions/list", methods=["GET"])
def list_all_sessions():
    """List all sessions with user-friendly details for dropdown selection"""
    try:
        # Get all sessions sorted by creation date (newest first)
        try:
            # Try MongoDB-style sorting first
            sessions = list(sessions_collection.find({}).sort("created_at", -1))
        except:
            # Fallback for local storage - get all sessions and sort in Python
            sessions = list(sessions_collection.find({}))
            # Sort by created_at if it exists, handling both datetime and string formats
            sessions.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        session_list = []
        for session in sessions:
            session_id = session["session_id"]
            
            # Count documents for this session
            documents_path = get_documents_path(session_id)
            documents_count = 0
            if os.path.exists(documents_path):
                documents_count = len([f for f in os.listdir(documents_path) 
                                     if os.path.isfile(os.path.join(documents_path, f))])
            
            # Check vector store status
            vector_store_path = get_vector_store_path(session_id)
            vector_store_ready = os.path.exists(vector_store_path) and len(os.listdir(vector_store_path)) > 0
            
            # Get recent activity (last message timestamp)
            try:
                last_message = conversations_collection.find_one(
                    {"session_id": session_id},
                    sort=[("timestamp", -1)]
                )
                last_activity = last_message["timestamp"] if last_message else session.get("created_at")
            except:
                # Fallback for local storage
                messages = list(conversations_collection.find({"session_id": session_id}))
                if messages:
                    # Sort messages by timestamp
                    messages.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
                    last_activity = messages[0]["timestamp"]
                else:
                    last_activity = session.get("created_at")
            
            # Create user-friendly display name
            display_name = session.get("user_description", "Unnamed Bot")
            if not display_name or display_name.strip() == "":
                display_name = f"{session.get('use_case', 'General')} Bot"
            
            # Handle datetime formatting
            created_at_str = ""
            if session.get("created_at"):
                if hasattr(session["created_at"], 'isoformat'):
                    created_at_str = session["created_at"].isoformat()
                else:
                    created_at_str = str(session["created_at"])
            
            last_activity_str = ""
            if last_activity:
                if hasattr(last_activity, 'isoformat'):
                    last_activity_str = last_activity.isoformat()
                else:
                    last_activity_str = str(last_activity)
            
            session_info = {
                "session_id": session_id,
                "display_name": display_name,
                "use_case": session.get("use_case", "General"),
                "documents_count": documents_count,
                "vector_store_ready": vector_store_ready,
                "has_custom_prompt": bool(session.get("custom_prompt", "").strip()),
                "created_at": created_at_str,
                "last_activity": last_activity_str,
                "short_id": session_id[:8]  # First 8 characters for easy reference
            }
            session_list.append(session_info)
        
        return jsonify({
            "sessions": session_list,
            "total_count": len(session_list)
        })
    
    except Exception as e:
        print(f"Error in list_all_sessions: {e}")  # Debug print
        return jsonify({"error": f"Failed to list sessions: {str(e)}"}), 500

@app.route("/api/sessions/<session_id>/status", methods=["GET"])
def get_session_status(session_id):
    """Get session status and configuration"""
    session_data = sessions_collection.find_one({"session_id": session_id})
    
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    # Check if vector store exists
    vector_store_path = get_vector_store_path(session_id)
    vector_store_ready = os.path.exists(vector_store_path) and os.listdir(vector_store_path)
    
    # Count documents
    documents_path = get_documents_path(session_id)
    documents_count = len([f for f in os.listdir(documents_path) 
                          if os.path.isfile(os.path.join(documents_path, f))])
    
    # Handle datetime formatting for both MongoDB and local storage
    created_at_str = ""
    if session_data.get("created_at"):
        created_at = session_data["created_at"]
        if hasattr(created_at, 'isoformat'):
            # It's a datetime object (MongoDB)
            created_at_str = created_at.isoformat()
        else:
            # It's already a string (local storage)
            created_at_str = str(created_at)
    
    return jsonify({
        "session_id": session_id,
        "user_description": session_data.get("user_description", ""),
        "use_case": session_data.get("use_case", ""),
        "documents_count": documents_count,
        "custom_prompt": session_data.get("custom_prompt", ""),
        "vector_store_ready": vector_store_ready,
        "created_at": created_at_str
    })

@app.route("/api/sessions/<session_id>/documents/upload", methods=["POST"])
def upload_documents(session_id):
    """Upload and process training documents"""
    # Verify session exists
    session_data = sessions_collection.find_one({"session_id": session_id})
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    if 'files' not in request.files:
        return jsonify({"error": "No files provided"}), 400
    
    files = request.files.getlist('files')
    uploaded_files = []
    errors = []
    
    documents_path = get_documents_path(session_id)
    
    for file in files:
        if file.filename == '':
            continue
            
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(documents_path, filename)
            
            try:
                file.save(file_path)
                uploaded_files.append(filename)
            except Exception as e:
                errors.append(f"Error saving {filename}: {str(e)}")
        else:
            errors.append(f"File type not allowed: {file.filename}")
    
    # Rebuild vector store
    vector_store_updated = False
    if uploaded_files:
        vector_store_updated = build_vector_store_for_session(session_id)
        
        # Update documents count in database
        documents_count = len([f for f in os.listdir(documents_path) 
                              if os.path.isfile(os.path.join(documents_path, f))])
        sessions_collection.update_one(
            {"session_id": session_id},
            {"$set": {"documents_count": documents_count}}
        )
    
    return jsonify({
        "uploaded_files": uploaded_files,
        "errors": errors,
        "vector_store_updated": vector_store_updated,
        "processing_status": "completed" if vector_store_updated else "failed"
    })

@app.route("/api/sessions/<session_id>/documents", methods=["GET"])
def list_documents(session_id):
    """List all uploaded documents for a session"""
    # Verify session exists
    session_data = sessions_collection.find_one({"session_id": session_id})
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    documents_path = get_documents_path(session_id)
    documents = []
    
    if os.path.exists(documents_path):
        for filename in os.listdir(documents_path):
            file_path = os.path.join(documents_path, filename)
            if os.path.isfile(file_path):
                stat = os.stat(file_path)
                documents.append({
                    "filename": filename,
                    "size": stat.st_size,
                    "upload_date": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "status": "processed"
                })
    
    return jsonify({"documents": documents})

@app.route("/api/sessions/<session_id>/documents/<filename>", methods=["DELETE"])
def delete_document(session_id, filename):
    """Remove a specific document and rebuild vector store"""
    # Verify session exists
    session_data = sessions_collection.find_one({"session_id": session_id})
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    documents_path = get_documents_path(session_id)
    file_path = os.path.join(documents_path, secure_filename(filename))
    
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found"}), 404
    
    try:
        os.remove(file_path)
        
        # Rebuild vector store
        vector_store_updated = build_vector_store_for_session(session_id)
        
        # Update documents count
        documents_count = len([f for f in os.listdir(documents_path) 
                              if os.path.isfile(os.path.join(documents_path, f))])
        sessions_collection.update_one(
            {"session_id": session_id},
            {"$set": {"documents_count": documents_count}}
        )
        
        return jsonify({
            "deleted": filename,
            "vector_store_updated": vector_store_updated
        })
    except Exception as e:
        return jsonify({"error": f"Error deleting file: {str(e)}"}), 500

@app.route("/api/sessions/<session_id>/prompt", methods=["PUT"])
def update_prompt(session_id):
    """Set or update custom instruction prompt"""
    # Verify session exists
    session_data = sessions_collection.find_one({"session_id": session_id})
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    data = request.get_json()
    if not data or "custom_prompt" not in data:
        return jsonify({"error": "custom_prompt is required"}), 400
    
    custom_prompt = data["custom_prompt"]
    
    # Update session in database
    sessions_collection.update_one(
        {"session_id": session_id},
        {"$set": {"custom_prompt": custom_prompt}}
    )
    
    return jsonify({
        "prompt_updated": True,
        "session_id": session_id
    })

@app.route("/api/sessions/<session_id>/prompt", methods=["GET"])
def get_prompt(session_id):
    """Retrieve current custom prompt"""
    session_data = sessions_collection.find_one({"session_id": session_id})
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    return jsonify({
        "custom_prompt": session_data.get("custom_prompt", ""),
        "default_prompt": "You are a helpful AI assistant."
    })

@app.route("/api/sessions/<session_id>/chat", methods=["POST"])
def chat_with_session(session_id):
    """Send a message and receive a response"""
    # Verify session exists
    session_data = sessions_collection.find_one({"session_id": session_id})
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    data = request.get_json()
    if not data or "message" not in data:
        return jsonify({"error": "message is required"}), 400
    
    user_message = data["message"]
    conversation_id = data.get("conversation_id", str(uuid.uuid4()))
    
    # Save user message
    conversations_collection.insert_one({
        "session_id": session_id,
        "conversation_id": conversation_id,
        "message": user_message,
        "message_type": "user",
        "timestamp": datetime.utcnow()
    })
    
    # Get bot response
    bot_response = query_llm_with_session(session_id, user_message, conversation_id)
    
    # Save bot response
    conversations_collection.insert_one({
        "session_id": session_id,
        "conversation_id": conversation_id,
        "message": bot_response,
        "message_type": "bot",
        "timestamp": datetime.utcnow()
    })
    
    # Check if context was used
    context_used = bool(retrieve_context_for_session(session_id, user_message))
    
    return jsonify({
        "response": bot_response,
        "conversation_id": conversation_id,
        "context_used": context_used
    })

@app.route("/api/sessions/<session_id>/conversations", methods=["GET"])
def get_conversations(session_id):
    """Retrieve conversation history"""
    # Verify session exists
    session_data = sessions_collection.find_one({"session_id": session_id})
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    # Get all conversations for this session
    messages = list(conversations_collection.find(
        {"session_id": session_id},
        {"_id": 0}
    ).sort("timestamp", 1))
    
    # Group by conversation_id
    conversations = {}
    for msg in messages:
        conv_id = msg["conversation_id"]
        if conv_id not in conversations:
            conversations[conv_id] = []
        conversations[conv_id].append({
            "message": msg["message"],
            "message_type": msg["message_type"],
            "timestamp": msg["timestamp"].isoformat()
        })
    
    # Convert to list format
    conversation_list = []
    for conv_id, msgs in conversations.items():
        conversation_list.append({
            "id": conv_id,
            "messages": msgs
        })
    
    return jsonify({"conversations": conversation_list})

@app.route("/api/sessions/<session_id>/conversations/<conversation_id>", methods=["DELETE"])
def clear_conversation(session_id, conversation_id):
    """Clear specific conversation history"""
    # Verify session exists
    session_data = sessions_collection.find_one({"session_id": session_id})
    if not session_data:
        return jsonify({"error": "Session not found"}), 404
    
    # Delete conversation messages
    result = conversations_collection.delete_many({
        "session_id": session_id,
        "conversation_id": conversation_id
    })
    
    return jsonify({
        "conversation_cleared": True,
        "messages_deleted": result.deleted_count
    })

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5002, debug=True)

